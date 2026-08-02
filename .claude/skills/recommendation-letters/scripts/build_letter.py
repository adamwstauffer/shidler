#!/usr/bin/env python3
"""Render a recommendation letter onto Adam W. Stauffer's branded letterhead base.

The branded base (`recommendations/recommendation-template-format.docx`, GITIGNORED because it
embeds Adam's signature image) supplies the UH/Shidler letterhead header, footers, and the
signature block (with signature + seal images). This script replaces only the BODY — title, date,
recipient block, salutation, and body paragraphs — and leaves everything from "Sincerely," onward
(the signature block and its images) untouched.

Usage:
    python build_letter.py content.json

content.json:
    {
      "base": "C:/GitHub/shidler/recommendations/recommendation-template-format.docx",
      "out":  "C:/.../Recommendation Letter - Name.docx",
      "title": "RECOMMENDATION LETTER:  Full Name",
      "date":  "July 31, 2026",
      "recipient_lines": ["Admissions Committee", "Program", "Institution"],  # [] for a generic letter
      "salutation": "Dear Members of the Admissions Committee,",              # or "To Whom It May Concern,"
      "body_paragraphs": ["para 1 ...", "para 2 ...", ...],
      "space_after": 240   # optional twips (default 240); lower to 120–160 to keep to one page
    }

Keep letters to ONE page unless there is a valid reason for more (see SKILL.md). If a letter spills
to a second page, re-run with a smaller "space_after" (e.g. 160 or 120) before shortening content.
"""
from __future__ import annotations
import json, sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Times New Roman"
SZ_HALFPTS = "24"  # 12pt


def _style_run(run, bold: bool) -> None:
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), FONT)
    for tag in ("w:sz", "w:szCs"):
        e = rpr.find(qn(tag))
        if e is None:
            e = rpr.makeelement(qn(tag), {})
            rpr.append(e)
        e.set(qn("w:val"), SZ_HALFPTS)


def _add(anchor, text, *, bold=False, center=False, before=0, after=0, line=1.0):
    p = anchor.insert_paragraph_before(text)
    _style_run(p.runs[0], bold)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(before / 20.0)
    pf.space_after = Pt(after / 20.0)
    pf.line_spacing = line
    return p


def build(cfg: dict) -> str:
    doc = Document(cfg["base"])
    # locate the signature block start ("Sincerely,") — everything before it is old body
    anchor = next((p for p in doc.paragraphs if p.text.strip().lower().startswith("sincerely")), None)
    if anchor is None:
        raise SystemExit("could not find the 'Sincerely,' paragraph in the base template")
    for p in list(doc.paragraphs):
        if p._element is anchor._element:  # compare XML elements, not fresh wrapper objects
            break
        p._element.getparent().remove(p._element)

    # Single-gap spacing (before=0, after=N) keeps consecutive paragraphs from stacking
    # before+after gaps — the main lever for keeping a letter to one page.
    after = int(cfg.get("space_after", 160))
    ls = float(cfg.get("line_spacing", 1.0))  # nudge to ~0.9 to keep a long letter to one page
    # title_before clears the letterhead image in the first-page header (~500+ twips)
    _add(anchor, cfg["title"], bold=True, center=True, before=int(cfg.get("title_before", 560)), after=after)
    _add(anchor, cfg["date"], before=0, after=after, line=ls)
    for line in cfg.get("recipient_lines", []):
        _add(anchor, line, before=0, after=0, line=ls)
    _add(anchor, cfg["salutation"], before=0, after=after, line=ls)
    for para in cfg["body_paragraphs"]:
        _add(anchor, para, before=0, after=after, line=ls)

    doc.save(cfg["out"])
    return cfg["out"]


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(__doc__)
        raise SystemExit(1)
    with open(sys.argv[1], encoding="utf-8") as fh:
        out = build(json.load(fh))
    print("wrote", out)

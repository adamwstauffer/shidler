"""FIN-321 Stage 3 (post-build FX hedging spec) grading scanner.

Walks the Stage 3 submissions directory, dedupes by student ID (latest
timestamp), inspects each spec deliverable (.md / .docx / .pdf / .txt), and
produces a grading worksheet with rubric signals and tentative auto-scores.

Mirrors `grade_stage2.py` in folder/dedup logic and curve-formula shape, but
the rubric is the four-criterion spec rubric defined in
`project-fx-hedging/stage3-spec-assignment.md`:

    Clarity & Professionalism   /1
    Analytical Logic            /1
    Completeness                /1
    Reproducibility             /1
                                = /4 total

The curve uses an 80% floor (per instructor direction): the lowest curved
score never drops below 0.80 * 4 = 3.20.
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

STAGE3_DIR = Path(
    r"C:\GitHub\shidler\courses\International-Finance-And-Securities\FIN-321"
    r"\ignore\2026-Spring\stage3"
)
OUTPUT_PATH = STAGE3_DIR / "_grading" / "stage3-grading-worksheet.xlsx"

# Folder name pattern: "<id>-<course> - <Name> - <Mon D, YYYY HHMM AM/PM>"
FOLDER_RE = re.compile(
    r"^(?P<sid>\d+)-\d+\s*-\s*(?P<name>.+?)\s*-\s*"
    r"(?P<month>[A-Za-z]+)\s+(?P<day>\d+),\s*(?P<year>\d{4})\s+"
    r"(?P<h>\d{1,4})\s*(?P<ampm>AM|PM)\s*$"
)

# Required sections from stage3-spec-assignment.md
REQUIRED_SECTIONS = [
    ("problem_statement", ["problem statement", "objective", "scope", "exposure"]),
    ("inputs",            ["inputs", "known variables"]),
    ("assumptions",       ["assumptions", "constraints"]),
    ("calculation_flow",  ["calculation flow", "calc flow", "calculation logic",
                           "workflow", "step-by-step", "step by step"]),
    ("outputs",           ["outputs"]),
    ("model_review",      ["model review", "what worked", "what to improve"]),
    ("sensitivity_plan",  ["sensitivity plan", "sensitivity analysis", "sensitivity table"]),
    ("limitations",       ["limitations", "next steps"]),
]

# Four core hedge strategies + sensitivity (mirrors stage2 grader's keywords)
HEDGE_KEYWORDS = {
    "Forward":     ["forward hedge", "forward rate", "forward contract",
                    "locked-in", "locked in", "usd_forward"],
    "MoneyMarket": ["money market", "mm hedge", "money-market", "borrow eur",
                    "borrow fc", "borrow foreign", "synthetic forward",
                    "covered interest", "parity"],
    "Put":         ["put option", "put hedge", "put premium", "k_put",
                    "put strike", "eur put", "premium_put", "prem_put",
                    "usd_put"],
    "Call":        ["call option", "call hedge", "call premium", "k_call",
                    "call strike", "eur call", "premium_call", "prem_call",
                    "usd_call"],
    "Sensitivity": ["sensitivity", "scenario table", "s_t", "sₜ",
                    "±5", "+/-5", "±5%", "scenarios"],
}

# Named-range tokens — both the canonical skeleton names and common alts.
# We match them as case-insensitive whole-token substrings in the spec text.
NAMED_RANGE_TOKENS = [
    # Canonical skeleton (from assignment table)
    "FC_AMT", "S0_in", "F0_in", "R_USD", "R_FC",
    "K_PUT", "K_CALL", "PREM_PUT", "PREM_CALL", "T_DAYS",
    # Template variants (lower-case / subscript)
    "S0", "F0", "r_USD", "r_EUR", "K_put", "K_call",
    "Premium_put", "Premium_call",
    # Output ranges suggested in template
    "USD_forward", "USD_mm", "USD_put", "USD_call",
    # Common alt tokens observed in stage2 grader
    "call_price", "call_strike", "put_price", "put_strike",
    "forward_price", "current_spot_price", "future_spot_price",
    "contract_notional", "rate_us", "rate_uk", "receivable",
    "payable", "scenario", "notional",
]
# Compile a single big alternation, case-insensitive, with word boundaries
_NAMED_RE = re.compile(
    r"(?<![A-Za-z0-9_])(?:" + "|".join(re.escape(t) for t in NAMED_RANGE_TOKENS)
    + r")(?![A-Za-z0-9_])",
    re.IGNORECASE,
)


@dataclass
class Submission:
    student_id: str
    student_name: str
    submitted_at: datetime
    folder: Path
    spec_file: Path | None = None


@dataclass
class Grade:
    student_id: str
    student_name: str
    submitted_at: datetime
    spec_filename: str
    file_format: str = ""
    word_count: int = 0
    table_count: int = 0
    code_block_count: int = 0
    sections_found: list[str] = field(default_factory=list)
    has_metadata_header: bool = False
    hedges_found: list[str] = field(default_factory=list)
    distinct_named_tokens: int = 0
    named_token_examples: list[str] = field(default_factory=list)
    model_review_words: int = 0
    auto_clarity: int = 0           # /1
    auto_logic: int = 0             # /1
    auto_completeness: int = 0      # /1
    auto_reproducibility: int = 0   # /1
    flags: list[str] = field(default_factory=list)
    error: str = ""


# ------------------------------------------------------------------
# Folder parsing & file discovery
# ------------------------------------------------------------------

def parse_folder(folder: Path) -> Submission | None:
    m = FOLDER_RE.match(folder.name)
    if not m:
        return None
    h = m.group("h")
    if len(h) == 3:
        hour, minute = int(h[0]), int(h[1:])
    elif len(h) == 4:
        hour, minute = int(h[:2]), int(h[2:])
    else:
        hour, minute = int(h), 0
    ampm = m.group("ampm").upper()
    if ampm == "PM" and hour != 12:
        hour += 12
    elif ampm == "AM" and hour == 12:
        hour = 0
    try:
        dt = datetime.strptime(
            f"{m.group('month')} {m.group('day')} {m.group('year')}", "%b %d %Y"
        ).replace(hour=hour, minute=minute)
    except ValueError:
        return None
    return Submission(
        student_id=m.group("sid"),
        student_name=m.group("name").strip(),
        submitted_at=dt,
        folder=folder,
    )


def find_spec_file(folder: Path) -> Path | None:
    """Pick best deliverable: prefer .md > .docx > .pdf > .txt."""
    by_ext: dict[str, list[Path]] = {".md": [], ".docx": [], ".pdf": [], ".txt": []}
    for p in folder.iterdir():
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext in by_ext:
            by_ext[ext].append(p)
    for ext in (".md", ".docx", ".pdf", ".txt"):
        if by_ext[ext]:
            by_ext[ext].sort(key=lambda p: len(p.name))
            return by_ext[ext][0]
    return None


def collect_submissions() -> list[Submission]:
    by_id: dict[str, Submission] = {}
    skipped = []
    for child in STAGE3_DIR.iterdir():
        if not child.is_dir() or child.name.startswith("_"):
            continue
        sub = parse_folder(child)
        if sub is None:
            skipped.append(child.name)
            continue
        existing = by_id.get(sub.student_id)
        if existing is None or sub.submitted_at > existing.submitted_at:
            by_id[sub.student_id] = sub
    if skipped:
        print(f"[warn] skipped {len(skipped)} unparsable folders:", skipped)
    subs = sorted(by_id.values(), key=lambda s: s.student_name.lower())
    for s in subs:
        s.spec_file = find_spec_file(s.folder)
    return subs


# ------------------------------------------------------------------
# Text extraction
# ------------------------------------------------------------------

def _extract_md_or_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pdfminer.high_level import extract_text
    try:
        return extract_text(str(path)) or ""
    except Exception as e:
        return f"[pdf extraction failed: {e}]"


def extract_text(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext == ".md":
        return _extract_md_or_txt(path), "md"
    if ext == ".txt":
        return _extract_md_or_txt(path), "txt"
    if ext == ".docx":
        return _extract_docx(path), "docx"
    if ext == ".pdf":
        return _extract_pdf(path), "pdf"
    return "", "unknown"


# ------------------------------------------------------------------
# Inspection / scoring
# ------------------------------------------------------------------

WORD_RE = re.compile(r"\b\w+\b")


def _count_words(text: str) -> int:
    return len(WORD_RE.findall(text))


def _count_md_tables(text: str) -> int:
    return len(re.findall(r"\n\s*\|[^\n]*\|\s*\n\s*\|[\s\-:|]+\|\s*\n", text))


def _count_code_blocks(text: str) -> int:
    return len(re.findall(r"```", text)) // 2


def _has_metadata_header(text: str) -> bool:
    head = text[:2000].lower()
    hits = sum(1 for tag in (
        "role:", "audience:", "purpose:", "version:", "created by:",
        "updated by:", "date created:", "date updated:",
    ) if tag in head)
    return hits >= 2


def _section_words_after(text: str, anchors: list[str], limit: int = 4000) -> int:
    """Word count from anchor heading to next H1/H2 (or end)."""
    lower = text.lower()
    best_idx = -1
    for a in anchors:
        i = lower.find(a)
        if i != -1 and (best_idx == -1 or i < best_idx):
            best_idx = i
    if best_idx == -1:
        return 0
    nl = text.find("\n", best_idx)
    start = nl + 1 if nl != -1 else best_idx
    tail = text[start:start + limit]
    end = re.search(r"\n#{1,2}\s+", tail)
    if end:
        tail = tail[: end.start()]
    return _count_words(tail)


def inspect(sub: Submission) -> Grade:
    g = Grade(
        student_id=sub.student_id,
        student_name=sub.student_name,
        submitted_at=sub.submitted_at,
        spec_filename=sub.spec_file.name if sub.spec_file else "(none)",
    )
    if sub.spec_file is None:
        g.error = "no spec file in folder"
        g.flags.append("NO_FILE")
        return g

    try:
        text, fmt = extract_text(sub.spec_file)
    except Exception as e:
        g.error = f"extract failed: {e}"
        g.flags.append("EXTRACT_FAILED")
        return g
    g.file_format = fmt

    if not text.strip():
        g.error = "empty spec"
        g.flags.append("EMPTY")
        return g

    # Strip markdown backslash-escapes before underscores so e.g. K\_PUT
    # is treated the same as K_PUT.
    text = text.replace("\\_", "_")
    lower = text.lower()
    g.word_count = _count_words(text)
    g.table_count = _count_md_tables(text)
    g.code_block_count = _count_code_blocks(text)
    g.has_metadata_header = _has_metadata_header(text)

    # Sections
    found_sections = []
    for key, anchors in REQUIRED_SECTIONS:
        if any(a in lower for a in anchors):
            found_sections.append(key)
    g.sections_found = found_sections

    # Hedge categories
    hedges = []
    for cat, keywords in HEDGE_KEYWORDS.items():
        if any(k in lower for k in keywords):
            hedges.append(cat)
    g.hedges_found = hedges

    # Named ranges
    matches = _NAMED_RE.findall(text)
    distinct = sorted({m.lower() for m in matches})
    g.distinct_named_tokens = len(distinct)
    g.named_token_examples = distinct[:8]

    # Model review depth
    g.model_review_words = _section_words_after(
        text, ["model review", "what worked", "what to improve"], limit=2500
    )

    # ----- Auto-scoring -----

    # 1. Clarity & Professionalism /1: any 2 of 3 (metadata, length, tables/code)
    clarity_signals = sum([
        g.has_metadata_header,
        g.word_count >= 800,
        (g.table_count >= 1) or (g.code_block_count >= 1),
    ])
    g.auto_clarity = 1 if clarity_signals >= 2 else 0

    # 2. Analytical Logic /1:
    #    >=3 of 4 core hedges (Forward/MM/Put/Call) AND calculation_flow section
    core_hedges = {"Forward", "MoneyMarket", "Put", "Call"}
    core_hits = len(core_hedges & set(hedges))
    has_calc = "calculation_flow" in found_sections
    g.auto_logic = 1 if (core_hits >= 3 and has_calc) else 0

    # 3. Completeness /1:
    #    >=6 of 8 sections AND model review present + >=100 words
    has_review = "model_review" in found_sections
    g.auto_completeness = 1 if (
        len(found_sections) >= 6 and has_review and g.model_review_words >= 100
    ) else 0

    # 4. Reproducibility /1:
    #    >=6 distinct named-range tokens (out of ~10 canonical inputs)
    g.auto_reproducibility = 1 if g.distinct_named_tokens >= 6 else 0

    # ----- Flags -----
    if g.word_count < 100:
        g.flags.append("STUB")
    missing_sections = [k for k, _ in REQUIRED_SECTIONS if k not in found_sections]
    if missing_sections:
        g.flags.append("MISSING_SECTION:" + ",".join(missing_sections))
    if core_hits < 3:
        g.flags.append("MISSING_HEDGE")
    if g.distinct_named_tokens < 6:
        g.flags.append("FEW_NAMED_RANGES")
    if g.word_count < 800:
        g.flags.append("SHORT")
    if g.model_review_words < 100:
        g.flags.append("THIN_MODEL_REVIEW")
    if "Sensitivity" not in hedges and "sensitivity_plan" not in found_sections:
        g.flags.append("NO_SENSITIVITY")

    return g


# ------------------------------------------------------------------
# Worksheet builder
# ------------------------------------------------------------------

LETTER_GRADE_SCALE = [
    ("A+", 97, None), ("A", 93, 97), ("A-", 90, 93),
    ("B+", 87, 90),  ("B", 83, 87), ("B-", 80, 83),
    ("C+", 77, 80),  ("C", 73, 77), ("C-", 70, 73),
    ("D+", 67, 70),  ("D", 65, 67), ("F", 0, 65),
]

TOTAL_POINTS = 4


def _write_curve_formulas(ws, final_col: int, rank_col: int,
                          quart_col: int, curved_col: int, generous_col: int,
                          first_row: int, last_row: int) -> None:
    """Quartile-graded curve, never reduces a raw score, rounds up to 0.05.

    Curved /4 (80% floor): top -> bottom
        Q1: 4.00 -> 3.75 -> Q2: 3.50 -> Q3: 3.35 -> Q4: 3.20  (= 80% of 4)
    Curved /4 (90% floor): generous variant
        Q1: 4.00 -> 3.90 -> 3.80 -> 3.70 -> 3.60                (= 90% of 4)
    """
    fl = get_column_letter(final_col)
    rl = get_column_letter(rank_col)
    rng = f"${fl}${first_row}:${fl}${last_row}"
    n_expr = f'COUNTIF({rng},">0")'

    curved_fill = PatternFill("solid", fgColor="E2EFDA")
    generous_fill = PatternFill("solid", fgColor="C6E0B4")
    for r in range(first_row, last_row + 1):
        final_ref = f"{fl}{r}"
        rank_ref = f"{rl}{r}"

        rank_f = (
            f'=IF({final_ref}=0,"",'
            f'RANK({final_ref},{rng},0)'
            f'+SUMPRODUCT(({rng}={final_ref})*(ROW({rng})<ROW())))'
        )
        ws.cell(row=r, column=rank_col, value=rank_f)
        quart_f = (
            f'=IF({final_ref}=0,"",'
            f'IF({rank_ref}<={n_expr}/4,1,'
            f'IF({rank_ref}<={n_expr}/2,2,'
            f'IF({rank_ref}<={n_expr}*3/4,3,4))))'
        )
        ws.cell(row=r, column=quart_col, value=quart_f)

        p = f'(({rank_ref}-1)/({n_expr}-1))'

        # 80% floor: 4.00 -> 3.75 -> 3.50 -> 3.35 -> 3.20
        floor80 = (
            f'IF({p}<=0.25,4-{p}*1,'
            f'IF({p}<=0.5,3.75-({p}-0.25)*1,'
            f'IF({p}<=0.75,3.5-({p}-0.5)*0.6,'
            f'3.35-({p}-0.75)*0.6)))'
        )
        ws.cell(
            row=r, column=curved_col,
            value=f'=IF({final_ref}=0,0,CEILING(MAX({final_ref},{floor80}),0.05))'
        ).fill = curved_fill

        # 90% floor: 4.00 -> 3.90 -> 3.80 -> 3.70 -> 3.60
        floor90 = (
            f'IF({p}<=0.25,4-{p}*0.4,'
            f'IF({p}<=0.5,3.9-({p}-0.25)*0.4,'
            f'IF({p}<=0.75,3.8-({p}-0.5)*0.4,'
            f'3.7-({p}-0.75)*0.4)))'
        )
        ws.cell(
            row=r, column=generous_col,
            value=f'=IF({final_ref}=0,0,CEILING(MAX({final_ref},{floor90}),0.05))'
        ).fill = generous_fill


def _write_letter_grade_summary(sm, curved_col: int, generous_col: int,
                                first_row: int, last_row: int) -> None:
    cv_l = get_column_letter(curved_col)
    gn_l = get_column_letter(generous_col)
    cv_rng = f"Grading!${cv_l}${first_row}:${cv_l}${last_row}"
    gn_rng = f"Grading!${gn_l}${first_row}:${gn_l}${last_row}"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="024731")

    sm.append([])
    headers = ["Letter", "Min %", f"Min /{TOTAL_POINTS}",
               "Count (Curved 80%)", "Count (Curved 90%)",
               "Histogram (80%)", "Histogram (90%)"]
    sm.append(headers)
    hr = sm.max_row
    for col_idx in range(1, len(headers) + 1):
        c = sm.cell(row=hr, column=col_idx)
        c.font = header_font
        c.fill = header_fill

    for letter, min_pct, max_pct in LETTER_GRADE_SCALE:
        min_pts = round(min_pct * TOTAL_POINTS / 100, 4)
        r = sm.max_row + 1
        if letter == "F":
            max_pts = round(max_pct * TOTAL_POINTS / 100, 4)
            cv_f = f'=COUNTIFS({cv_rng},">0",{cv_rng},"<"&{max_pts})'
            gn_f = f'=COUNTIFS({gn_rng},">0",{gn_rng},"<"&{max_pts})'
        elif max_pct is None:
            cv_f = f'=COUNTIF({cv_rng},">="&{min_pts})'
            gn_f = f'=COUNTIF({gn_rng},">="&{min_pts})'
        else:
            max_pts = round(max_pct * TOTAL_POINTS / 100, 4)
            cv_f = f'=COUNTIFS({cv_rng},">="&{min_pts},{cv_rng},"<"&{max_pts})'
            gn_f = f'=COUNTIFS({gn_rng},">="&{min_pts},{gn_rng},"<"&{max_pts})'
        sm.cell(row=r, column=1, value=letter)
        sm.cell(row=r, column=2, value=min_pct)
        sm.cell(row=r, column=3, value=min_pts)
        sm.cell(row=r, column=4, value=cv_f)
        sm.cell(row=r, column=5, value=gn_f)
        sm.cell(row=r, column=6, value=f'=REPT("█",D{r})')
        sm.cell(row=r, column=7, value=f'=REPT("█",E{r})')

    r = sm.max_row + 1
    sm.cell(row=r, column=1, value="No submission")
    sm.cell(row=r, column=4, value=f'=COUNTIF({cv_rng},0)')
    sm.cell(row=r, column=5, value=f'=COUNTIF({gn_rng},0)')
    sm.cell(row=r, column=6, value=f'=REPT("█",D{r})')
    sm.cell(row=r, column=7, value=f'=REPT("█",E{r})')

    sm.column_dimensions["F"].width = 40
    sm.column_dimensions["G"].width = 40


def build_worksheet(grades: list[Grade]) -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Grading"

    headers = [
        "Student ID", "Student Name", "Submitted", "File", "Format",
        "Word Count", "# Tables", "# Code Blocks", "Metadata Header",
        "Sections Found", "# Sections",
        "Hedges Found", "# Hedges",
        "# Distinct Named Tokens", "Named Token Examples",
        "Model Review Words",
        "Auto Clarity /1", "Auto Logic /1",
        "Auto Completeness /1", "Auto Reproducibility /1",
        "Auto Total /4", "Final /4",
        "Rank", "Quartile", "Curved /4 (80%)", "Curved /4 (90%)",
        "Flags", "Comments",
    ]
    ws.append(headers)

    FINAL_COL = headers.index("Final /4") + 1
    RANK_COL = headers.index("Rank") + 1
    QUART_COL = headers.index("Quartile") + 1
    CURVED_COL = headers.index("Curved /4 (80%)") + 1
    GENEROUS_COL = headers.index("Curved /4 (90%)") + 1
    FLAGS_COL = headers.index("Flags") + 1

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="024731")
    for col, _ in enumerate(headers, 1):
        c = ws.cell(row=1, column=col)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(vertical="center", wrap_text=True)

    flag_fill = PatternFill("solid", fgColor="FFF2CC")
    error_fill = PatternFill("solid", fgColor="F8CBAD")

    for g in grades:
        auto_total = (
            g.auto_clarity + g.auto_logic
            + g.auto_completeness + g.auto_reproducibility
        )
        row = [
            g.student_id,
            g.student_name,
            g.submitted_at.strftime("%Y-%m-%d %H:%M"),
            g.spec_filename,
            g.file_format,
            g.word_count,
            g.table_count,
            g.code_block_count,
            "Y" if g.has_metadata_header else "",
            ", ".join(g.sections_found),
            len(g.sections_found),
            ", ".join(g.hedges_found),
            len(g.hedges_found),
            g.distinct_named_tokens,
            ", ".join(g.named_token_examples),
            g.model_review_words,
            g.auto_clarity, g.auto_logic,
            g.auto_completeness, g.auto_reproducibility,
            auto_total, auto_total,  # Final defaults to Auto Total for editing
            None, None, None, None,  # Rank, Quartile, Curved, Curved 90% formulas
            ", ".join(g.flags),
            g.error,
        ]
        ws.append(row)
        r = ws.max_row
        if g.error:
            for col in range(1, len(headers) + 1):
                ws.cell(row=r, column=col).fill = error_fill
        elif g.flags:
            ws.cell(row=r, column=FLAGS_COL).fill = flag_fill

    _write_curve_formulas(ws, FINAL_COL, RANK_COL, QUART_COL, CURVED_COL,
                          GENEROUS_COL, 2, ws.max_row)

    widths = [
        11, 30, 17, 36, 8,
        11, 9, 11, 12,
        38, 11,
        30, 10,
        12, 30,
        14,
        12, 12, 14, 16,
        13, 10,
        8, 10, 14, 14,
        38, 28,
    ]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "F2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"

    # Summary sheet
    sm = wb.create_sheet("Summary")
    sm.append(["Metric", "Value"])
    sm.append(["Total unique students", len(grades)])
    sm.append(["Missing file", sum(1 for g in grades if "NO_FILE" in g.flags)])
    sm.append(["Stub submission", sum(1 for g in grades if "STUB" in g.flags)])
    sm.append(["Missing section(s)", sum(1 for g in grades
                                         if any(f.startswith("MISSING_SECTION") for f in g.flags))])
    sm.append(["Missing hedge type", sum(1 for g in grades if "MISSING_HEDGE" in g.flags)])
    sm.append(["Few named ranges", sum(1 for g in grades if "FEW_NAMED_RANGES" in g.flags)])
    sm.append(["Short (<800 words)", sum(1 for g in grades if "SHORT" in g.flags)])
    sm.append(["Thin model review", sum(1 for g in grades if "THIN_MODEL_REVIEW" in g.flags)])
    sm.append(["No sensitivity", sum(1 for g in grades if "NO_SENSITIVITY" in g.flags)])
    if grades:
        avg = sum(
            g.auto_clarity + g.auto_logic + g.auto_completeness + g.auto_reproducibility
            for g in grades
        ) / len(grades)
        sm.append(["Average auto-score /4", round(avg, 2)])

    sm.append([])
    sm.append(["Curve policy", ""])
    sm.append(["Ranking basis", "Final /4 (updates live)"])
    sm.append(["Mode",
               "Ceiling-floor + round up: Curved = CEILING(MAX(Final, quartile floor), 0.05). "
               "Rounds up to nearest 0.05; never reduces raw score."])
    sm.append(["Tiebreak", "Sheet order (alphabetical by name)"])
    sm.append(["Curved /4 (80% floor)", ""])
    sm.append(["Q1 top 25%", "4.00 -> 3.75"])
    sm.append(["Q2", "3.75 -> 3.50"])
    sm.append(["Q3", "3.50 -> 3.35"])
    sm.append(["Q4 bottom 25%", "3.35 -> 3.20  (= 80% of 4)"])
    sm.append(["Final = 0", "Curved = 0 (non-submission)"])
    sm.append([])
    sm.append(["Curved /4 (90% floor) — generous option", ""])
    sm.append(["Q1", "4.00 -> 3.90"])
    sm.append(["Q2", "3.90 -> 3.80"])
    sm.append(["Q3", "3.80 -> 3.70"])
    sm.append(["Q4", "3.70 -> 3.60  (= 90% of 4)"])

    for col in (1, 2):
        sm.column_dimensions[get_column_letter(col)].width = 30
    for cell in sm[1]:
        cell.font = header_font
        cell.fill = header_fill

    _write_letter_grade_summary(sm, CURVED_COL, GENEROUS_COL, 2, ws.max_row)

    wb.save(OUTPUT_PATH)


def main() -> int:
    subs = collect_submissions()
    print(f"Found {len(subs)} unique students (deduped by ID).")
    grades: list[Grade] = []
    for s in subs:
        print(f"  grading {s.student_id} {s.student_name} ...", end=" ")
        g = inspect(s)
        grades.append(g)
        if g.error:
            print(f"ERROR: {g.error}")
        else:
            total = g.auto_clarity + g.auto_logic + g.auto_completeness + g.auto_reproducibility
            print(f"auto={total}/4 flags={','.join(g.flags) or '-'}")
    build_worksheet(grades)
    print(f"\nWrote {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())

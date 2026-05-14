"""BUS-314 Stage 4 (final analysis + structured AI prompt) grading scanner.

Walks the Stage 4 submissions directory, dedupes by student ID (keeping the
latest timestamp per ID), inspects each deliverable (.md / .docx / .pdf), and
produces a grading worksheet with rubric signals and tentative auto-scores.

Mirrors grade_stage3.py in shape, but the rubric is the six-criterion Stage 4
rubric defined in `accounting-ratios/stage4-final-analysis-assignment.md`,
collapsed to one point per criterion for a /6 total:

    Ratio Interpretation                 /1
    Strategic Recommendations            /1
    Du Pont Analysis                     /1
    Structured AI Prompt                 /1
    Professionalism & Communication      /1
    AI Output OR Manual Analysis         /1
                                         = /6 total

Plus a flat -0.5 deduction when no GitHub link / reference is present
(per instructor direction).

The curve uses an 85% floor: the lowest curved score never drops below
0.85 * 6 = 5.10.
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

STAGE4_DIR = Path(
    r"C:\GitHub\shidler\courses\BUS-314-International-Corporate-Finance"
    r"\ignore-term\2026 Spring\stage4"
)
OUTPUT_PATH = STAGE4_DIR / "_grading" / "stage4-grading-worksheet.xlsx"
INDEX_HTML = STAGE4_DIR / "index.html"

# Folder name pattern: "<id>-<course> - <Name> - <Mon D, YYYY HHMM AM/PM>"
FOLDER_RE = re.compile(
    r"^(?P<sid>\d+)-\d+\s*-\s*(?P<name>.+?)\s*-\s*"
    r"(?P<month>[A-Za-z]+)\s+(?P<day>\d+),\s*(?P<year>\d{4})\s+"
    r"(?P<h>\d{1,4})\s*(?P<ampm>AM|PM)\s*$"
)

# Six ratio categories — coverage is the basis for "Ratio Interpretation"
CATEGORY_KEYWORDS = {
    "Performance":   ["market value added", "mva", "market-to-book",
                      "market to book", "economic value added", "eva"],
    "Profitability": ["roa", "roc", "roe", "return on assets",
                      "return on capital", "return on equity"],
    "Efficiency":    ["asset turnover", "receivables turnover",
                      "collection period", "inventory turnover",
                      "days in inventory", "profit margin"],
    "Leverage":      ["debt ratio", "debt-equity", "debt to equity",
                      "times interest", "cash coverage", "debt burden",
                      "leverage"],
    "Liquidity":     ["nwc", "current ratio", "quick ratio", "cash ratio",
                      "working capital"],
    "Du Pont":       ["du pont", "dupont"],
}

# Named-range conventions — same prefixes as Stage 2/3
PREFIX_PATTERNS = {
    "BAL_":   re.compile(r"\bBAL_[A-Za-z0-9_]+", re.IGNORECASE),
    "INC_":   re.compile(r"\bINC_[A-Za-z0-9_]+", re.IGNORECASE),
    "CASH_":  re.compile(r"\bCASH_[A-Za-z0-9_]+", re.IGNORECASE),
    "RATIO_": re.compile(r"\bRATIO_[A-Za-z0-9_]+", re.IGNORECASE),
    "scope_": re.compile(r"\b(?:startYear|currentYear|avg)_[A-Za-z0-9_]+",
                         re.IGNORECASE),
}

# AI-tool / LLM mentions (signal of AI-generated output OR explicit AI tooling)
LLM_MENTION_RE = re.compile(
    r"\b(claude|chatgpt|gpt[-\s]?\d|openai|gemini|copilot|llm|"
    r"large language model|anthropic|llama)\b",
    re.IGNORECASE,
)

# Recommendation-section detector — match as a heading-like line ending with
# "Recommendation(s)". Must be a short line (heading-style), not a paragraph
# containing the word. Tolerates markdown heading markers (`##`), bolding
# (`**`), and bare lines from docx/pdf extraction. Uses `[ \t]*` rather than
# `\s*` so the match doesn't span backward across blank lines.
RECOMMENDATION_HEADER_RE = re.compile(
    r"(?im)^[ \t]*(?:#{1,6}[ \t]+|\*\*[ \t]*)?[^\n]{0,55}?"
    r"recommendations?"
    r"[ \t]*\*{0,2}[ \t]*:?[ \t]*$",
)

# Recommendation-table header detector — for docx/pdf where the recs section
# became a flattened table like "# | Recommendation | Supporting Data | ...".
RECOMMENDATION_TABLE_HEADER_RE = re.compile(
    r"(?im)^[^\n]{0,40}\brecommendation\b[^\n]{0,8}\|", )

# Du Pont / decomposition language
DUPONT_DRIVERS_RE = re.compile(
    r"\b(margin|turnover|leverage|debt burden)\b", re.IGNORECASE
)

# Structured-prompt signal — section headers used inside the prompt body
PROMPT_SECTION_HEADERS = [
    "# goal", "# company financial data", "# ratio formulas", "# verification",
    "# workbook structure", "# color coding", "# named range",
]

# GitHub link / reference detector
GITHUB_RE = re.compile(
    r"(github\.com|github\.io|raw\.githubusercontent|githubusercontent\.com|"
    r"\bgithub\b)",
    re.IGNORECASE,
)


@dataclass
class Submission:
    student_id: str
    student_name: str
    submitted_at: datetime
    folder: Path
    deliverable: Path | None = None
    comments: str = ""


@dataclass
class Grade:
    student_id: str
    student_name: str
    submitted_at: datetime
    filename: str
    file_format: str = ""
    word_count: int = 0
    table_count: int = 0
    code_block_count: int = 0
    has_metadata_header: bool = False
    categories_found: list[str] = field(default_factory=list)
    named_range_prefixes_hit: list[str] = field(default_factory=list)
    distinct_named_tokens: int = 0
    has_recommendations: bool = False
    recommendation_count: int = 0
    has_dupont: bool = False
    dupont_driver_hits: int = 0
    has_prompt_block: bool = False
    prompt_section_hits: int = 0
    llm_mentioned: bool = False
    has_github: bool = False
    github_hits: int = 0
    github_in_comments: bool = False
    auto_interpretation: int = 0   # /1
    auto_recommendations: int = 0  # /1
    auto_dupont: int = 0           # /1
    auto_prompt: int = 0           # /1
    auto_professionalism: int = 0  # /1
    auto_ai_or_manual: int = 0     # /1
    github_deduction: float = 0.0  # 0 or 0.5
    github_bonus: float = 0.0      # 0 or 0.25
    flags: list[str] = field(default_factory=list)
    error: str = ""


# ------------------------------------------------------------------
# Folder parsing & file discovery
# ------------------------------------------------------------------

def parse_folder(folder: Path) -> Submission | None:
    m = FOLDER_RE.match(folder.name)
    if not m:
        return None
    h = m.group("h")
    if len(h) == 3:
        hour, minute = int(h[0]), int(h[1:])
    elif len(h) == 4:
        hour, minute = int(h[:2]), int(h[2:])
    else:
        hour, minute = int(h), 0
    ampm = m.group("ampm").upper()
    if ampm == "PM" and hour != 12:
        hour += 12
    elif ampm == "AM" and hour == 12:
        hour = 0
    try:
        dt = datetime.strptime(
            f"{m.group('month')} {m.group('day')} {m.group('year')}", "%b %d %Y"
        ).replace(hour=hour, minute=minute)
    except ValueError:
        return None
    return Submission(
        student_id=m.group("sid"),
        student_name=m.group("name").strip(),
        submitted_at=dt,
        folder=folder,
    )


def find_deliverable(folder: Path) -> Path | None:
    """Pick the best deliverable: prefer .md over .docx over .pdf.

    If multiple files of the same format exist, pick the shortest filename
    (matches Stage 2/3 heuristic — usually the canonical submission).
    """
    by_ext: dict[str, list[Path]] = {".md": [], ".docx": [], ".pdf": []}
    for p in folder.iterdir():
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext in by_ext:
            by_ext[ext].append(p)
    for ext in (".md", ".docx", ".pdf"):
        if by_ext[ext]:
            by_ext[ext].sort(key=lambda p: len(p.name))
            return by_ext[ext][0]
    return None


def parse_index_comments(path: Path) -> dict[str, str]:
    """Parse the LMS index.html and return {normalized_name: comments_text}.

    The HTML format is one giant table; each student appears as:
        <b>{Last}, {First}</b>  ... <b>Comments:</b>{html}</td>
    Comments often contain GitHub links that are not present in the deliverable
    file itself, so we surface them for the GitHub-reference check.
    """
    if not path.exists():
        return {}
    html = path.read_text(encoding="utf-8", errors="ignore")
    out: dict[str, str] = {}
    pattern = re.compile(
        r"<b>([^<,]+),\s*([^<]+?)</b>"
        r".*?<b>Comments:</b>(.*?)</td>",
        re.DOTALL | re.IGNORECASE,
    )
    for m in pattern.finditer(html):
        last = m.group(1).strip()
        first = m.group(2).strip()
        # Folder names are "Last First" — match that ordering with last
        # potentially containing spaces (e.g., "Mena Rieke Claudia").
        name = f"{last} {first}".lower()
        # Strip HTML tags from the comments block, keep href URLs.
        text = m.group(3)
        # Pull href URLs out so they survive the tag strip.
        text = re.sub(r'<a[^>]*href="([^"]+)"[^>]*>',
                      lambda mo: " " + mo.group(1) + " ", text,
                      flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"&nbsp;", " ", text)
        out[name] = text.strip()
    return out


def collect_submissions() -> list[Submission]:
    by_id: dict[str, Submission] = {}
    skipped = []
    for child in STAGE4_DIR.iterdir():
        if not child.is_dir() or child.name.startswith("_"):
            continue
        sub = parse_folder(child)
        if sub is None:
            skipped.append(child.name)
            continue
        existing = by_id.get(sub.student_id)
        if existing is None or sub.submitted_at > existing.submitted_at:
            by_id[sub.student_id] = sub
    if skipped:
        print(f"[warn] skipped {len(skipped)} unparsable folders:", skipped)
    subs = sorted(by_id.values(), key=lambda s: s.student_name.lower())
    comments_by_name = parse_index_comments(INDEX_HTML)
    for s in subs:
        s.deliverable = find_deliverable(s.folder)
        s.comments = comments_by_name.get(s.student_name.lower(), "")
    missing = [s.student_name for s in subs if not s.comments]
    if missing and comments_by_name:
        print(f"[info] no LMS comments matched for {len(missing)} students: {missing}")
    return subs


# ------------------------------------------------------------------
# Text extraction
# ------------------------------------------------------------------

def _extract_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pdfminer.high_level import extract_text
    try:
        return extract_text(str(path)) or ""
    except Exception as e:
        return f"[pdf extraction failed: {e}]"


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, format_label)."""
    ext = path.suffix.lower()
    if ext == ".md":
        return _extract_md(path), "md"
    if ext == ".docx":
        return _extract_docx(path), "docx"
    if ext == ".pdf":
        return _extract_pdf(path), "pdf"
    return "", "unknown"


# ------------------------------------------------------------------
# Inspection / scoring
# ------------------------------------------------------------------

WORD_RE = re.compile(r"\b\w+\b")


def _count_words(text: str) -> int:
    return len(WORD_RE.findall(text))


def _count_md_tables(text: str) -> int:
    """Count markdown tables (header row + separator-row pairs)."""
    return len(re.findall(r"\n\s*\|[^\n]*\|\s*\n\s*\|[\s\-:|]+\|\s*\n", text))


def _count_code_blocks(text: str) -> int:
    return len(re.findall(r"```", text)) // 2


def _has_metadata_header(text: str) -> bool:
    """Look for a header block with role/audience/date-style metadata."""
    head = text[:2000].lower()
    hits = sum(1 for tag in (
        "role:", "audience:", "purpose:", "version:", "prepared by:",
        "created by:", "updated by:", "date:", "author:", "course:",
    ) if tag in head)
    return hits >= 2


def _count_recommendations(text: str) -> int:
    """Heuristic: count numbered/bulleted items in any 'Recommendations' section.

    The header must look like a markdown heading or bolded label — not a
    substring of the document title. We scan all matches and pick the one
    whose section yields the highest item count, since some authors place
    multiple "Recommendation" labels (e.g., one in the toc, one as a heading).
    """
    best = 0

    # Approach 1: scope to a "Recommendations" section heading and count items.
    # Number-followed-by separator: period, paren, colon, hyphen, en/em-dash.
    sep_class = r"[\.\):\-–—]"
    for m in RECOMMENDATION_HEADER_RE.finditer(text):
        # Re-anchor: jump past the matched header line so the section bound
        # search doesn't see the header as its own delimiter.
        idx = m.end()
        nl = text.find("\n", idx)
        start = nl + 1 if nl != -1 else idx
        tail = text[start:start + 6000]
        end = re.search(r"\n#{1,2}[ \t]+", tail)
        if end:
            tail = tail[: end.start()]
        # Numbered list items at line start: "1.", "**1.", "1)"
        numbered = len(re.findall(
            r"(?m)^[ \t]*(?:\*\*|__)?[ \t]*\d+" + sep_class + r"[ \t]+", tail))
        # Numbered H3/H4-style headings: "### 1. Title" or "### 1 — Title"
        numbered_heading = len(re.findall(
            r"(?m)^[ \t]*#{2,6}[ \t]+\d+[ \t]*" + sep_class + r"[ \t]+", tail))
        # Heading-style "### Recommendation 1:" / "### Recommendation 2:"
        recommendation_heading = len(re.findall(
            r"(?im)^[ \t]*#{2,6}[ \t]+recommendation[ \t]*\d+", tail))
        # Plain bulleted items
        bulleted = len(re.findall(r"(?m)^[ \t]*[-*+][ \t]+\S", tail))
        # Table data rows ("| 1 |" or "| 1. |" or bare "1 |")
        table_numbered = len(re.findall(
            r"(?m)^[ \t]*\|?[ \t]*\d+[\.\)]?[ \t]*\|", tail))
        best = max(best, numbered, numbered_heading,
                   recommendation_heading, bulleted, table_numbered)

    # Approach 2: docx/pdf flattened table — find a "Recommendation | ..." row
    # and count the numbered data rows that follow it.
    for m in RECOMMENDATION_TABLE_HEADER_RE.finditer(text):
        nl = text.find("\n", m.start())
        if nl == -1:
            continue
        tail = text[nl + 1: nl + 1 + 4000]
        end = re.search(r"\n#{1,2}[ \t]+", tail)
        if end:
            tail = tail[: end.start()]
        rows = len(re.findall(r"(?m)^[ \t]*\|?[ \t]*\d+[ \t]*\|", tail))
        best = max(best, rows)

    return best


def inspect(sub: Submission) -> Grade:
    g = Grade(
        student_id=sub.student_id,
        student_name=sub.student_name,
        submitted_at=sub.submitted_at,
        filename=sub.deliverable.name if sub.deliverable else "(none)",
    )
    if sub.deliverable is None:
        g.error = "no deliverable in folder"
        g.flags.append("NO_FILE")
        return g

    try:
        text, fmt = extract_text(sub.deliverable)
    except Exception as e:
        g.error = f"extract failed: {e}"
        g.flags.append("EXTRACT_FAILED")
        return g
    g.file_format = fmt

    if not text.strip():
        g.error = "empty deliverable"
        g.flags.append("EMPTY")
        return g

    # Strip markdown backslash-escapes before underscores so e.g. BAL\_total
    # is treated the same as BAL_total when matching named ranges.
    text = text.replace("\\_", "_")
    lower = text.lower()
    g.word_count = _count_words(text)
    g.table_count = _count_md_tables(text)
    g.code_block_count = _count_code_blocks(text)
    g.has_metadata_header = _has_metadata_header(text)

    # Categories
    cats_found = []
    for cat, keywords in CATEGORY_KEYWORDS.items():
        if any(k in lower for k in keywords):
            cats_found.append(cat)
    g.categories_found = cats_found

    # Named-range prefixes
    prefixes_hit = []
    distinct_tokens: set[str] = set()
    for label, pat in PREFIX_PATTERNS.items():
        matches = pat.findall(text)
        if matches:
            prefixes_hit.append(label)
            for m in matches:
                distinct_tokens.add(m.lower())
    g.named_range_prefixes_hit = prefixes_hit
    g.distinct_named_tokens = len(distinct_tokens)

    # Recommendations
    g.recommendation_count = _count_recommendations(text)
    g.has_recommendations = g.recommendation_count >= 2

    # Du Pont
    g.has_dupont = "du pont" in lower or "dupont" in lower
    g.dupont_driver_hits = len(set(
        m.lower() for m in DUPONT_DRIVERS_RE.findall(text)
    ))

    # Structured-prompt signal
    g.has_prompt_block = g.code_block_count >= 1 or "structured ai prompt" in lower
    g.prompt_section_hits = sum(1 for h in PROMPT_SECTION_HEADERS if h in lower)

    # LLM / AI mention
    g.llm_mentioned = bool(LLM_MENTION_RE.search(text))

    # GitHub link / reference — check both the deliverable text and the LMS
    # comments block (where students often paste their repo URL).
    github_in_text = GITHUB_RE.findall(text)
    github_in_comments = GITHUB_RE.findall(sub.comments) if sub.comments else []
    g.github_hits = len(github_in_text) + len(github_in_comments)
    g.github_in_comments = bool(github_in_comments)
    g.has_github = g.github_hits > 0
    g.github_deduction = 0.0 if g.has_github else 0.5
    g.github_bonus = 0.25 if g.has_github else 0.0

    # ----- Auto-scoring (each /1) -----

    # 1. Ratio Interpretation /1
    #    >=4 of 6 categories AND substantive analysis (>= 600 words OR >=1 table)
    g.auto_interpretation = 1 if (
        len(cats_found) >= 4 and (g.word_count >= 600 or g.table_count >= 1)
    ) else 0

    # 2. Strategic Recommendations /1
    #    Recommendations section present AND >= 2 distinct items
    g.auto_recommendations = 1 if g.has_recommendations else 0

    # 3. Du Pont Analysis /1
    #    Du Pont mentioned AND >= 2 driver concepts referenced
    g.auto_dupont = 1 if (g.has_dupont and g.dupont_driver_hits >= 2) else 0

    # 4. Structured AI Prompt /1
    #    Prompt block present AND (>= 2 prompt section headers OR >= 8 named-range tokens)
    g.auto_prompt = 1 if (
        g.has_prompt_block
        and (g.prompt_section_hits >= 2 or g.distinct_named_tokens >= 8)
    ) else 0

    # 5. Professionalism & Communication /1
    #    metadata header + length >= 800 words + (>=1 table OR >=1 code block)
    prof_signals = sum([
        g.has_metadata_header,
        g.word_count >= 800,
        (g.table_count >= 1) or (g.code_block_count >= 1),
    ])
    g.auto_professionalism = 1 if prof_signals >= 2 else 0

    # 6. AI Output OR Manual Analysis /1
    #    Either LLM tool referenced, OR substantive manual analysis
    #    (>= 5 categories AND >= 1000 words)
    manual_substantive = len(cats_found) >= 5 and g.word_count >= 1000
    g.auto_ai_or_manual = 1 if (g.llm_mentioned or manual_substantive) else 0

    # ----- Flags -----
    if g.word_count < 50:
        g.flags.append("STUB")
    if not g.has_github:
        g.flags.append("NO_GITHUB")
    if len(cats_found) < 4:
        g.flags.append("MISSING_CATEGORY")
    if not g.has_dupont:
        g.flags.append("NO_DUPONT")
    if not g.has_recommendations:
        g.flags.append("THIN_RECOMMENDATIONS")
    if not g.has_prompt_block:
        g.flags.append("NO_PROMPT_BLOCK")
    if g.distinct_named_tokens < 8:
        g.flags.append("FEW_NAMED_RANGES")
    if g.word_count < 800:
        g.flags.append("SHORT")

    return g


# ------------------------------------------------------------------
# Worksheet builder
# ------------------------------------------------------------------

LETTER_GRADE_SCALE = [
    ("A+", 97, None), ("A", 93, 97), ("A-", 90, 93),
    ("B+", 87, 90),  ("B", 83, 87), ("B-", 80, 83),
    ("C+", 77, 80),  ("C", 73, 77), ("C-", 70, 73),
    ("D+", 67, 70),  ("D", 65, 67), ("F", 0, 65),
]

TOTAL_POINTS = 6  # Stage 4 rubric is /6


def _write_curve_formulas(ws, final_col: int, rank_col: int,
                          quart_col: int, curved_col: int, generous_col: int,
                          first_row: int, last_row: int) -> None:
    """Quartile-graded curve, never reduces a raw score, rounds up to 0.05.

    Curved /6 (85% floor): top -> bottom
        Q1: 6.00 -> 5.70  (drops 0.30)
        Q2: 5.70 -> 5.40  (drops 0.30)
        Q3: 5.40 -> 5.25  (drops 0.15)
        Q4: 5.25 -> 5.10  (drops 0.15)
        Bottom of curve = 5.10 = 85% of 6.

    Curved 90% /6 (generous variant):
        Q1: 6.00 -> 5.85
        Q2: 5.85 -> 5.70
        Q3: 5.70 -> 5.55
        Q4: 5.55 -> 5.40
        Bottom = 5.40 = 90% of 6.
    """
    fl = get_column_letter(final_col)
    rl = get_column_letter(rank_col)
    rng = f"${fl}${first_row}:${fl}${last_row}"
    n_expr = f'COUNTIF({rng},">0")'

    curved_fill = PatternFill("solid", fgColor="E2EFDA")
    generous_fill = PatternFill("solid", fgColor="C6E0B4")
    for r in range(first_row, last_row + 1):
        final_ref = f"{fl}{r}"
        rank_ref = f"{rl}{r}"

        rank_f = (
            f'=IF({final_ref}=0,"",'
            f'RANK({final_ref},{rng},0)'
            f'+SUMPRODUCT(({rng}={final_ref})*(ROW({rng})<ROW())))'
        )
        ws.cell(row=r, column=rank_col, value=rank_f)
        quart_f = (
            f'=IF({final_ref}=0,"",'
            f'IF({rank_ref}<={n_expr}/4,1,'
            f'IF({rank_ref}<={n_expr}/2,2,'
            f'IF({rank_ref}<={n_expr}*3/4,3,4))))'
        )
        ws.cell(row=r, column=quart_col, value=quart_f)

        p = f'(({rank_ref}-1)/({n_expr}-1))'

        # 85% floor curve: 6.00 -> 5.70 -> 5.40 -> 5.25 -> 5.10
        floor85 = (
            f'IF({p}<=0.25,6-{p}*1.2,'
            f'IF({p}<=0.5,5.7-({p}-0.25)*1.2,'
            f'IF({p}<=0.75,5.4-({p}-0.5)*0.6,'
            f'5.25-({p}-0.75)*0.6)))'
        )
        ws.cell(
            row=r, column=curved_col,
            value=f'=IF({final_ref}=0,0,CEILING(MAX({final_ref},{floor85}),0.05))'
        ).fill = curved_fill

        # 90% floor curve (generous): 6.00 -> 5.85 -> 5.70 -> 5.55 -> 5.40
        floor90 = (
            f'IF({p}<=0.25,6-{p}*0.6,'
            f'IF({p}<=0.5,5.85-({p}-0.25)*0.6,'
            f'IF({p}<=0.75,5.7-({p}-0.5)*0.6,'
            f'5.55-({p}-0.75)*0.6)))'
        )
        ws.cell(
            row=r, column=generous_col,
            value=f'=IF({final_ref}=0,0,CEILING(MAX({final_ref},{floor90}),0.05))'
        ).fill = generous_fill


def _write_letter_grade_summary(sm, curved_col: int, generous_col: int,
                                first_row: int, last_row: int) -> None:
    cv_l = get_column_letter(curved_col)
    gn_l = get_column_letter(generous_col)
    cv_rng = f"Grading!${cv_l}${first_row}:${cv_l}${last_row}"
    gn_rng = f"Grading!${gn_l}${first_row}:${gn_l}${last_row}"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="024731")

    sm.append([])
    headers = ["Letter", "Min %", f"Min /{TOTAL_POINTS}",
               "Count (Curved 85%)", "Count (Curved 90%)",
               "Histogram (85%)", "Histogram (90%)"]
    sm.append(headers)
    hr = sm.max_row
    for col_idx in range(1, len(headers) + 1):
        c = sm.cell(row=hr, column=col_idx)
        c.font = header_font
        c.fill = header_fill

    for letter, min_pct, max_pct in LETTER_GRADE_SCALE:
        min_pts = round(min_pct * TOTAL_POINTS / 100, 4)
        r = sm.max_row + 1
        if letter == "F":
            max_pts = round(max_pct * TOTAL_POINTS / 100, 4)
            cv_f = f'=COUNTIFS({cv_rng},">0",{cv_rng},"<"&{max_pts})'
            gn_f = f'=COUNTIFS({gn_rng},">0",{gn_rng},"<"&{max_pts})'
        elif max_pct is None:
            cv_f = f'=COUNTIF({cv_rng},">="&{min_pts})'
            gn_f = f'=COUNTIF({gn_rng},">="&{min_pts})'
        else:
            max_pts = round(max_pct * TOTAL_POINTS / 100, 4)
            cv_f = f'=COUNTIFS({cv_rng},">="&{min_pts},{cv_rng},"<"&{max_pts})'
            gn_f = f'=COUNTIFS({gn_rng},">="&{min_pts},{gn_rng},"<"&{max_pts})'
        sm.cell(row=r, column=1, value=letter)
        sm.cell(row=r, column=2, value=min_pct)
        sm.cell(row=r, column=3, value=min_pts)
        sm.cell(row=r, column=4, value=cv_f)
        sm.cell(row=r, column=5, value=gn_f)
        sm.cell(row=r, column=6, value=f'=REPT("█",D{r})')
        sm.cell(row=r, column=7, value=f'=REPT("█",E{r})')

    r = sm.max_row + 1
    sm.cell(row=r, column=1, value="No submission")
    sm.cell(row=r, column=4, value=f'=COUNTIF({cv_rng},0)')
    sm.cell(row=r, column=5, value=f'=COUNTIF({gn_rng},0)')
    sm.cell(row=r, column=6, value=f'=REPT("█",D{r})')
    sm.cell(row=r, column=7, value=f'=REPT("█",E{r})')

    sm.column_dimensions["F"].width = 40
    sm.column_dimensions["G"].width = 40


def build_worksheet(grades: list[Grade]) -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Grading"

    headers = [
        "Student ID", "Student Name", "Submitted", "File", "Format",
        "Word Count", "# Tables", "# Code Blocks", "Metadata Header",
        "Categories Found", "# Categories",
        "Named Prefix Families", "# Distinct Named Tokens",
        "# Recommendations", "Du Pont", "# Du Pont Drivers",
        "Prompt Block", "# Prompt Section Headers",
        "LLM Mentioned", "GitHub Found", "GitHub via Comments", "# GitHub Hits",
        "Auto Interpretation /1", "Auto Recommendations /1",
        "Auto Du Pont /1", "Auto Prompt /1",
        "Auto Professionalism /1", "Auto AI/Manual /1",
        "Auto Total /6", "GitHub Bonus", "GitHub Deduction", "Final /6",
        "Rank", "Quartile", "Curved /6 (85%)", "Curved /6 (90%)",
        "Flags", "Comments",
    ]
    ws.append(headers)

    FINAL_COL = headers.index("Final /6") + 1
    RANK_COL = headers.index("Rank") + 1
    QUART_COL = headers.index("Quartile") + 1
    CURVED_COL = headers.index("Curved /6 (85%)") + 1
    GENEROUS_COL = headers.index("Curved /6 (90%)") + 1
    FLAGS_COL = headers.index("Flags") + 1

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="024731")
    for col, _ in enumerate(headers, 1):
        c = ws.cell(row=1, column=col)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(vertical="center", wrap_text=True)

    flag_fill = PatternFill("solid", fgColor="FFF2CC")
    error_fill = PatternFill("solid", fgColor="F8CBAD")
    deduction_fill = PatternFill("solid", fgColor="FCE4D6")
    bonus_fill = PatternFill("solid", fgColor="E2EFDA")

    for g in grades:
        auto_total = (
            g.auto_interpretation + g.auto_recommendations
            + g.auto_dupont + g.auto_prompt
            + g.auto_professionalism + g.auto_ai_or_manual
        )
        final_score = max(0.0, auto_total + g.github_bonus - g.github_deduction)
        row = [
            g.student_id,
            g.student_name,
            g.submitted_at.strftime("%Y-%m-%d %H:%M"),
            g.filename,
            g.file_format,
            g.word_count,
            g.table_count,
            g.code_block_count,
            "Y" if g.has_metadata_header else "",
            ", ".join(g.categories_found),
            len(g.categories_found),
            ", ".join(g.named_range_prefixes_hit),
            g.distinct_named_tokens,
            g.recommendation_count,
            "Y" if g.has_dupont else "",
            g.dupont_driver_hits,
            "Y" if g.has_prompt_block else "",
            g.prompt_section_hits,
            "Y" if g.llm_mentioned else "",
            "Y" if g.has_github else "",
            "Y" if g.github_in_comments else "",
            g.github_hits,
            g.auto_interpretation, g.auto_recommendations,
            g.auto_dupont, g.auto_prompt,
            g.auto_professionalism, g.auto_ai_or_manual,
            auto_total, g.github_bonus, g.github_deduction, final_score,
            None, None, None, None,  # Rank, Quartile, Curved, Curved 90%
            ", ".join(g.flags),
            g.error,
        ]
        ws.append(row)
        r = ws.max_row
        if g.error:
            for col in range(1, len(headers) + 1):
                ws.cell(row=r, column=col).fill = error_fill
        else:
            if g.flags:
                ws.cell(row=r, column=FLAGS_COL).fill = flag_fill
            if g.github_deduction > 0:
                ws.cell(
                    row=r,
                    column=headers.index("GitHub Deduction") + 1,
                ).fill = deduction_fill
            if g.github_bonus > 0:
                ws.cell(
                    row=r,
                    column=headers.index("GitHub Bonus") + 1,
                ).fill = bonus_fill

    _write_curve_formulas(ws, FINAL_COL, RANK_COL, QUART_COL, CURVED_COL,
                          GENEROUS_COL, 2, ws.max_row)

    widths = [
        11, 24, 17, 36, 8,
        11, 9, 11, 12,
        30, 12,
        22, 12,
        14, 9, 14,
        12, 14,
        13, 13, 14, 13,
        16, 18,
        14, 12,
        18, 16,
        13, 14, 16, 11,
        8, 10, 14, 14,
        38, 28,
    ]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "F2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"

    # Summary sheet
    sm = wb.create_sheet("Summary")
    sm.append(["Metric", "Value"])
    sm.append(["Total unique students", len(grades)])
    sm.append(["Missing file", sum(1 for g in grades if "NO_FILE" in g.flags)])
    sm.append(["No GitHub link/reference",
               sum(1 for g in grades if "NO_GITHUB" in g.flags)])
    sm.append(["Missing ratio category",
               sum(1 for g in grades if "MISSING_CATEGORY" in g.flags)])
    sm.append(["No Du Pont", sum(1 for g in grades if "NO_DUPONT" in g.flags)])
    sm.append(["Thin recommendations",
               sum(1 for g in grades if "THIN_RECOMMENDATIONS" in g.flags)])
    sm.append(["No prompt block",
               sum(1 for g in grades if "NO_PROMPT_BLOCK" in g.flags)])
    sm.append(["Few named ranges",
               sum(1 for g in grades if "FEW_NAMED_RANGES" in g.flags)])
    sm.append(["Short (<800 words)",
               sum(1 for g in grades if "SHORT" in g.flags)])
    if grades:
        avg_auto = sum(
            g.auto_interpretation + g.auto_recommendations
            + g.auto_dupont + g.auto_prompt
            + g.auto_professionalism + g.auto_ai_or_manual
            for g in grades
        ) / len(grades)
        avg_final = sum(
            max(
                0.0,
                (g.auto_interpretation + g.auto_recommendations
                 + g.auto_dupont + g.auto_prompt
                 + g.auto_professionalism + g.auto_ai_or_manual)
                + g.github_bonus - g.github_deduction,
            )
            for g in grades
        ) / len(grades)
        sm.append(["Average auto-score /6 (pre-deduction)", round(avg_auto, 2)])
        sm.append(["Average final /6 (post-deduction)", round(avg_final, 2)])

    sm.append([])
    sm.append(["Curve policy", ""])
    sm.append(["Ranking basis", "Final /6 (updates live)"])
    sm.append(["Mode",
               "Ceiling-floor + round up: Curved = CEILING(MAX(Final, quartile floor), 0.05). "
               "Rounds up to nearest 0.05; never reduces raw score."])
    sm.append(["GitHub bonus", "+0.25 if github link/reference detected (in deliverable or LMS comments)"])
    sm.append(["GitHub deduction", "−0.5 if no github link/reference detected"])
    sm.append(["Tiebreak", "Sheet order (alphabetical by name)"])
    sm.append(["Curved /6 (85% floor)", ""])
    sm.append(["Q1 top 25%", "6.00 -> 5.70"])
    sm.append(["Q2", "5.70 -> 5.40"])
    sm.append(["Q3", "5.40 -> 5.25"])
    sm.append(["Q4 bottom 25%", "5.25 -> 5.10  (= 85% of 6)"])
    sm.append(["Final = 0", "Curved = 0 (non-submission)"])
    sm.append([])
    sm.append(["Curved /6 (90% floor) — generous option", ""])
    sm.append(["Q1", "6.00 -> 5.85"])
    sm.append(["Q2", "5.85 -> 5.70"])
    sm.append(["Q3", "5.70 -> 5.55"])
    sm.append(["Q4", "5.55 -> 5.40  (= 90% of 6)"])

    for col in (1, 2):
        sm.column_dimensions[get_column_letter(col)].width = 32
    for cell in sm[1]:
        cell.font = header_font
        cell.fill = header_fill

    _write_letter_grade_summary(sm, CURVED_COL, GENEROUS_COL, 2, ws.max_row)

    wb.save(OUTPUT_PATH)


def main() -> int:
    subs = collect_submissions()
    print(f"Found {len(subs)} unique students (deduped by ID).")
    grades: list[Grade] = []
    for s in subs:
        print(f"  grading {s.student_id} {s.student_name} ...", end=" ")
        g = inspect(s)
        grades.append(g)
        if g.error:
            print(f"ERROR: {g.error}")
        else:
            total = (
                g.auto_interpretation + g.auto_recommendations
                + g.auto_dupont + g.auto_prompt
                + g.auto_professionalism + g.auto_ai_or_manual
            )
            final = max(0.0, total + g.github_bonus - g.github_deduction)
            ded = f" -{g.github_deduction:.2f}" if g.github_deduction else ""
            bon = f" +{g.github_bonus:.2f}" if g.github_bonus else ""
            print(f"auto={total}/6{bon}{ded} -> {final}/6 flags={','.join(g.flags) or '-'}")
    build_worksheet(grades)
    print(f"\nWrote {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
